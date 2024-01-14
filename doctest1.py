
from docx import Document
from docx.enum.text import WD_UNDERLINE
import os
import requests
class Execute:
    '''
        Execute Paragraphs KeyWords Replace
        paragraph: docx paragraph
    '''

    def __init__(self, paragraph):
        self.paragraph = paragraph


    def p_replace(self, x:int, key:str, value:str):
        '''
        paragraph replace
        The reason why you do not replace the text in a paragraph directly is that it will cause the original format to
        change. Replacing the text in runs will not cause the original format to change
        :param x:       paragraph id
        :param key:     Keywords that need to be replaced
        :param value:   The replaced keywords
        :return:
        '''
        # Gets the coordinate index values of all the characters in this paragraph [{run_index , char_index}]
        p_maps = [{"run": y, "char": z} for y, run in enumerate(self.paragraph.runs) for z, char in enumerate(list(run.text))]
        # Handle the number of times key occurs in this paragraph, and record the starting position in the list.
        # Here, while self.text.find(key) >= 0, the {"ab":"abc"} term will enter an endless loop
        # Takes a single paragraph as an independent body and gets an index list of key positions within the paragraph, or if the paragraph contains multiple keys, there are multiple index values
        k_idx = [s for s in range(len(self.paragraph.text)) if self.paragraph.text.find(key, s, len(self.paragraph.text)) == s]
        for i, start_idx in enumerate(reversed(k_idx)):       # Reverse order iteration
            end_idx = start_idx + len(key)                    # The end position of the keyword in this paragraph
            k_maps = p_maps[start_idx:end_idx]                # Map Slice List A list of dictionaries for sections that contain keywords in a paragraph
            self.r_replace(k_maps, value)
            # print(f"\t |Paragraph {x+1: >3}, object {i+1: >3} replaced successfully! | {key} ===> {value}")


    def r_replace(self, k_maps:list, value:str):
        '''
        :param k_maps: The list of indexed dictionaries containing keywords， e.g:[{"run":15, "char":3},{"run":15, "char":4},{"run":16, "char":0}]
        :param value:
        :return:
        Accept arguments, removing the characters in k_maps from back to front, leaving the first one to replace with value
        Note: Must be removed in reverse order, otherwise the list length change will cause IndedxError: string index out of range
        '''
        for i, position in enumerate(reversed(k_maps), start=1):
            y, z = position["run"], position["char"]
            run:object = self.paragraph.runs[y]         # "k_maps" may contain multiple run ids, which need to be separated
            # Pit: Instead of the replace() method, str is converted to list after a single word to prevent run.text from making an error in some cases (e.g., a single run contains a duplicate word)
            thisrun = list(run.text)
            if i < len(k_maps):
                thisrun.pop(z)          # Deleting a corresponding word
            if i == len(k_maps):        # The last iteration (first word), that is, the number of iterations is equal to the length of k_maps
                thisrun[z] = value      # Replace the word in the corresponding position with the new content
            run.text = ''.join(thisrun) # Recover


class WordReplacer:
    def __init__(self, file):
        self.docx = Document(file)
    
    def replace_in_paragraph(self, paragraph, replace_dict):
        for idx, para in enumerate(self.docx.paragraphs):
            # if paragraph._element.xpath('.//w:endnoteReference'):
            #     print("This paragraph contains endnotes:")
            if para.text == paragraph:
                Execute(para).p_replace(idx, paragraph, replace_dict)
                print(replace_dict.strip())
                break

        for section in self.docx.sections:
            for header_paragraph in section.header.paragraphs:
                if header_paragraph.text == paragraph:
                    Execute(header_paragraph).p_replace(0, header_paragraph.text, replace_dict)
                    print(replace_dict.strip())

            for footer_paragraph in section.footer.paragraphs:
                if footer_paragraph.text == paragraph:
                    Execute(footer_paragraph).p_replace(0, footer_paragraph.text, replace_dict)
                    print(replace_dict.strip())
            for header_table in section.header.tables:
                for row in header_table.rows:
                    for cell in row.cells:
                        for cell_paragraph in cell.paragraphs:
                            if cell_paragraph.text == paragraph:
                                Execute(cell_paragraph).p_replace(0, cell_paragraph.text, replace_word)
                                print(replace_dict.strip())

            for footer_table in section.footer.tables:
                for row in footer_table.rows:
                    for cell in row.cells:
                        for cell_paragraph in cell.paragraphs:
                            if cell_paragraph.text == paragraph:
                                Execute(cell_paragraph).p_replace(0, cell_paragraph.text, replace_word)
                                print(replace_dict.strip())          

            
                                
    def replace_in_table(self, paragraph, replace_word):
        for table in self.docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for cell_paragraph in cell.paragraphs:
                        if cell_paragraph.text == paragraph:
                            Execute(cell_paragraph).p_replace(0, cell_paragraph.text, replace_word)
                            print(replace_word.strip())
        for section in self.docx.sections:
            for header_table in section.header.tables:
                for row in header_table.rows:
                    for cell in row.cells:
                        for cell_paragraph in cell.paragraphs:
                            if cell_paragraph.text == paragraph:
                                Execute(cell_paragraph).p_replace(0, cell_paragraph.text, replace_word)
                                print(replace_word.strip())

            for footer_table in section.footer.tables:
                for row in footer_table.rows:
                    for cell in row.cells:
                        for cell_paragraph in cell.paragraphs:
                            if cell_paragraph.text == paragraph:
                                Execute(cell_paragraph).p_replace(0, cell_paragraph.text, replace_word)
                                print(replace_word.strip())   

    def save(self, filepath:str):
        '''
        :param filepath: File saving path
        :return:
        '''
        print(filepath)
        self.docx.save(filepath)
        
    @staticmethod
    def docx_list(dirPath):
        '''
        :param dirPath:
        :return: List of docx files in the current directory
        '''
        fileList = []
        for roots, dirs, files in os.walk(dirPath):
            for file in files:
                if file.endswith("docx") and file[0] != "~":  # Find the docx document and exclude temporary files
                    fileRoot = os.path.join(roots, file)
                    fileList.append(fileRoot)
        print("This directory finds a total of {0} related files!".format(len(fileList)))
        return fileList

class WordUnderlineFinder:
    def is_underlined(self, run):
        """
        Check if the run is underlined.
        """
        return run.underline
    def collect_underlined_text(self, doc):
        """
        Collect underlined text from the document.
        """
        underlined_text = []
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if self.is_underlined(run):
                    underlined_text.append(run.text)
                    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for cell_paragraph in cell.paragraphs:
                        current_phrase = ""
                        for run in cell_paragraph.runs:
                            if self.is_underlined(run):
                                current_phrase += run.text
                        if (current_phrase != ""):
                            underlined_text.append(current_phrase)
        return underlined_text

def is_real_reference(ref_paragraph):
    """
    Check if "reference" is in real (non-bold and non-italic) text.
    """
    for run in ref_paragraph.runs:
        if "reference" in run.text.lower() and not (run.bold or run.italic):
            return True
            
    return False